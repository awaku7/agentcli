from __future__ import annotations

import json
from io import BytesIO
from pathlib import Path
from typing import Any

from .i18n_helper import make_tool_translator

_ = make_tool_translator(__file__)

try:
    from docx import Document
except Exception:  # pragma: no cover
    Document = None

try:
    from odf import teletype
    from odf.opendocument import load as odf_load
    from odf.table import Table, TableCell, TableRow
    from odf.text import H, P
except Exception:  # pragma: no cover
    teletype = None
    odf_load = None
    Table = TableCell = TableRow = H = P = None

try:
    from striprtf.striprtf import rtf_to_text
except Exception:  # pragma: no cover
    rtf_to_text = None

try:
    import msoffcrypto
except Exception:  # pragma: no cover
    msoffcrypto = None  # type: ignore[assignment]


TOOL_SPEC: dict[str, Any] = {
    "tool_level": 1,
    "tool_genre": "office",
    "type": "function",
    "function": {
        "name": "document_extract",
        "description": _(
            "tool.description",
            default=(
                "Extract text and basic structure from Word-like documents (.docx, .rtf, .odt). "
                "Choose output_format=text or output_format=json. The tool always returns JSON."
            ),
        ),
        "system_prompt": _(
            "tool.system_prompt",
            default=(
                "Read .docx, .rtf, or .odt documents and return JSON only. "
                "If output_format=text, place the extracted text in the JSON response. "
                "If output_format=json, include structured sections and tables when available."
            ),
        ),
        "x_search_terms": _(
            "x_search_terms",
            default=[
                "document_extract",
                "document extract",
                "docx",
                "rtf",
                "odt",
                "extract text",
            ],
        ),
        "x_search_terms_en": [
            "document_extract",
            "document extract",
            "docx",
            "rtf",
            "odt",
            "extract text",
        ],
        "parameters": {
            "type": "object",
            "properties": {
                "path": {
                    "type": "string",
                    "description": _(
                        "param.path.description",
                        default="Path to the input document (.docx, .rtf, .odt).",
                    ),
                },
                "password": {
                    "type": "string",
                    "description": _(
                        "param.password.description",
                        default=(
                            "Optional password for encrypted DOCX input. If omitted and the file is encrypted, "
                            "the tool will prompt once for a password."
                        ),
                    ),
                },
                "output_format": {
                    "type": "string",
                    "enum": ["text", "json"],
                    "default": "json",
                    "description": _(
                        "param.output_format.description",
                        default="Output format: text or json.",
                    ),
                },
            },
            "required": ["path"],
            "additionalProperties": False,
        },
    },
}

BUSY_LABEL = True


def _json_error(message: str, **extra: Any) -> dict[str, Any]:
    out: dict[str, Any] = {"ok": False, "error": message, "warnings": []}
    out.update(extra)
    return out


def _prompt_for_password(path: str) -> str | None:
    try:
        from .human_ask_tool import run_tool as human_ask
    except Exception:
        return None

    message = _(
        "prompt.password",
        default="Enter the password for this file:\n{path}",
    ).format(path=path)
    try:
        res_json = human_ask({"message": message, "is_password": True})
        res = json.loads(res_json)
    except Exception:
        return None

    pwd = str(res.get("user_reply") or "").strip()
    return pwd or None


def _docx_extract(path: Path, password: str | None = None) -> dict[str, Any]:
    if Document is None:
        raise RuntimeError("python-docx is not available")

    if msoffcrypto is not None:
        with path.open("rb") as fin:
            office = msoffcrypto.OfficeFile(fin)
            encrypted = False
            try:
                encrypted = bool(office.is_encrypted())
            except Exception:
                encrypted = False

            if encrypted:
                if not password:
                    password = _prompt_for_password(str(path))
                if not password:
                    raise RuntimeError("password is required for encrypted DOCX files")

                office.load_key(password=password)
                decrypted = BytesIO()
                office.decrypt(decrypted)
                decrypted.seek(0)
                doc = Document(decrypted)
            else:
                doc = Document(str(path))
    else:
        doc = Document(str(path))

    paragraphs: list[str] = []
    sections: list[dict[str, Any]] = []
    tables: list[list[list[str]]] = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        style_name = getattr(getattr(para, "style", None), "name", "") or ""
        level = None
        if style_name.lower().startswith("heading "):
            try:
                level = int(style_name.split()[-1])
            except Exception:
                level = None
        sections.append(
            {
                "type": "heading" if level is not None else "paragraph",
                "level": level,
                "text": text,
            }
        )
        paragraphs.append(text)

    for table in doc.tables:
        rows: list[list[str]] = []
        for row in table.rows:
            rows.append([cell.text.replace("\r", "\n").strip() for cell in row.cells])
        tables.append(rows)

    return {
        "file_type": "docx",
        "text": "\n\n".join(paragraphs),
        "sections": sections,
        "tables": tables,
    }


def _rtf_extract(path: Path) -> dict[str, Any]:
    if rtf_to_text is None:
        raise RuntimeError("striprtf is not available")

    raw = path.read_bytes()
    text = rtf_to_text(raw.decode("utf-8", errors="ignore"))
    lines = [line.strip() for line in text.splitlines() if line.strip()]

    return {
        "file_type": "rtf",
        "text": "\n".join(lines),
        "sections": [{"type": "paragraph", "text": line} for line in lines],
        "tables": [],
    }


def _odt_extract(path: Path) -> dict[str, Any]:
    if odf_load is None:
        raise RuntimeError("odfpy is not available")

    doc = odf_load(str(path))
    paragraphs: list[str] = []
    sections: list[dict[str, Any]] = []
    tables: list[list[list[str]]] = []

    for node in doc.getElementsByType(H):
        text = teletype.extractText(node) if teletype is not None else ""
        text = text.strip()
        if text:
            paragraphs.append(text)
            sections.append({"type": "heading", "level": None, "text": text})

    for node in doc.getElementsByType(P):
        text = teletype.extractText(node) if teletype is not None else ""
        text = text.strip()
        if text:
            paragraphs.append(text)
            sections.append({"type": "paragraph", "level": None, "text": text})

    for table in doc.getElementsByType(Table):
        rows: list[list[str]] = []
        for row in table.getElementsByType(TableRow):
            cells: list[str] = []
            for cell in row.getElementsByType(TableCell):
                cell_text = ""
                for para in cell.getElementsByType(P):
                    t = teletype.extractText(para) if teletype is not None else ""
                    if t:
                        cell_text += t
                cells.append(cell_text)
            rows.append(cells)
        tables.append(rows)

    return {
        "file_type": "odt",
        "text": "\n".join(paragraphs),
        "sections": sections,
        "tables": tables,
    }


def run(
    path: str, output_format: str = "json", password: str | None = None
) -> dict[str, Any]:
    p = Path(path)
    suffix = p.suffix.lower()
    if suffix not in {".docx", ".rtf", ".odt"}:
        return _json_error(f"unsupported file type: {suffix}")

    try:
        if suffix == ".docx":
            data = _docx_extract(p, password=password)
        elif suffix == ".rtf":
            data = _rtf_extract(p)
        else:
            data = _odt_extract(p)
    except Exception as exc:
        return _json_error(str(exc))

    out_format = (output_format or "json").strip().lower()
    if out_format not in {"text", "json"}:
        return _json_error(f"unsupported output_format: {out_format}")

    data = {**data, "ok": True, "warnings": [], "output_format": out_format}
    return data


def run_tool(args: dict[str, Any]) -> str:
    args = args or {}
    path = str(args.get("path") or "").strip()
    password = str(args.get("password") or "").strip() or None
    output_format = str(args.get("output_format") or "json").strip().lower()

    if not path:
        return json.dumps(_json_error("path is required"), ensure_ascii=False)

    result = run(path=path, output_format=output_format, password=password)
    return json.dumps(result, ensure_ascii=False)
