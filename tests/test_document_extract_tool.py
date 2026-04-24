from __future__ import annotations

import json
from pathlib import Path

import pytest

from uagent.tools.document_extract_tool import run_tool as document_extract


def _loads(out: str) -> dict:
    obj = json.loads(out)
    assert isinstance(obj, dict)
    return obj


def test_document_extract_rejects_unsupported_extension(repo_tmp_path: Path) -> None:
    path = repo_tmp_path / "sample.txt"
    path.write_text("hello", encoding="utf-8")

    out = _loads(document_extract({"path": str(path), "output_format": "json"}))
    assert out["ok"] is False
    assert "unsupported file type" in str(out["error"])


def test_document_extract_docx_smoke(repo_tmp_path: Path) -> None:
    pytest.importorskip("docx")
    from docx import Document as DocxDocument

    path = repo_tmp_path / "sample.docx"
    doc = DocxDocument()
    doc.add_heading("Title", level=1)
    doc.add_paragraph("Body")
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "A"
    table.cell(0, 1).text = "B"
    table.cell(1, 0).text = "1"
    table.cell(1, 1).text = "2"
    doc.save(str(path))

    out = _loads(document_extract({"path": str(path), "output_format": "json"}))
    assert out["ok"] is True
    assert out["file_type"] == "docx"
    assert out["output_format"] == "json"
    assert "Title" in str(out["text"])
    assert any(
        sec.get("type") == "heading"
        and sec.get("level") == 1
        and sec.get("text") == "Title"
        for sec in out["sections"]
    )
    assert out["tables"][0][0][0] == "A"


def test_document_extract_rtf_smoke(repo_tmp_path: Path) -> None:
    pytest.importorskip("striprtf")

    path = repo_tmp_path / "sample.rtf"
    path.write_text(r"{\rtf1\ansi Hello\par World\par}", encoding="utf-8")

    out = _loads(document_extract({"path": str(path), "output_format": "text"}))
    assert out["ok"] is True
    assert out["file_type"] == "rtf"
    assert out["output_format"] == "text"
    assert "Hello" in str(out["text"])
    assert "World" in str(out["text"])


def test_document_extract_odt_smoke(repo_tmp_path: Path) -> None:
    pytest.importorskip("odf")
    from odf.opendocument import OpenDocumentText
    from odf.text import H, P

    path = repo_tmp_path / "sample.odt"
    doc = OpenDocumentText()
    doc.text.addElement(H(outlinelevel=1, text="Title"))
    doc.text.addElement(P(text="Body"))
    doc.save(str(path))

    out = _loads(document_extract({"path": str(path), "output_format": "json"}))
    assert out["ok"] is True
    assert out["file_type"] == "odt"
    assert out["output_format"] == "json"
    assert "Title" in str(out["text"])
    assert "Body" in str(out["text"])
    assert out["sections"]
